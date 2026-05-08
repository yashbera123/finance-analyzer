"""
Transaction Parsing Engine
==========================
Heuristic-based parser that:
  1. Reads delimited text / spreadsheet files into a DataFrame
  2. Detects date, description, amount, and debit/credit columns
  3. Normalizes every row into a standard Transaction object
  4. Handles missing / dirty data gracefully

No AI calls — pure rule-based heuristics.
"""

from __future__ import annotations

import csv
import importlib.util
import logging
import re
import datetime as _dt
from datetime import date, datetime
from pathlib import Path
from typing import List, Optional, Tuple

import pandas as pd

from app.schemas.transaction import (
    ColumnMapping,
    ParseResult,
    ParseSummary,
    Transaction,
    TransactionType,
)

logger = logging.getLogger(__name__)

_TEXT_FILE_EXTENSIONS = {".csv", ".tsv", ".txt", ".pdf", ".docx", ".doc"}
_SPREADSHEET_EXTENSIONS = {".xlsx", ".xlsm", ".xls"}

# ═══════════════════════════════════════════════════════════════════════════
# Column-name heuristics (lowered, stripped)
# ═══════════════════════════════════════════════════════════════════════════

_DATE_KEYWORDS = [
    "date", "transaction date", "txn date", "trans date",
    "posting date", "value date", "trade date", "book date",
    "post date", "txn_date", "trans_date", "posting_date", "value_date",
    "post_date",
]

_DESCRIPTION_KEYWORDS = [
    "description", "narration", "details", "memo", "particulars",
    "remarks", "transaction description", "txn description",
    "payee", "merchant", "name", "reference", "note",
    "txn_description", "trans_description",
]

_AMOUNT_KEYWORDS = [
    "amount", "value", "sum", "total", "transaction amount",
    "txn amount", "net amount", "closing balance",
    "txn_amount", "trans_amount", "amt",
]

_DEBIT_KEYWORDS = [
    "debit", "withdrawal", "dr", "debit amount", "debit_amount",
    "withdrawals", "money out", "spent",
]

_CREDIT_KEYWORDS = [
    "credit", "deposit", "cr", "credit amount", "credit_amount",
    "deposits", "money in", "received",
]

_TYPE_KEYWORDS = [
    "type", "transaction type", "txn type", "dr/cr", "cr/dr",
    "direction", "flow",
]

_MIN_HEADER_HITS = 2

_HEADER_KEYWORDS = list(dict.fromkeys(
    _DATE_KEYWORDS
    + _DESCRIPTION_KEYWORDS
    + _AMOUNT_KEYWORDS
    + _DEBIT_KEYWORDS
    + _CREDIT_KEYWORDS
    + _TYPE_KEYWORDS
))


# ═══════════════════════════════════════════════════════════════════════════
# Date parsing — try many formats
# ═══════════════════════════════════════════════════════════════════════════

_DATE_FORMATS = [
    "%Y-%m-%d",       # 2024-01-15
    "%d-%m-%Y",       # 15-01-2024
    "%m-%d-%Y",       # 01-15-2024
    "%d/%m/%Y",       # 15/01/2024
    "%m/%d/%Y",       # 01/15/2024
    "%Y/%m/%d",       # 2024/01/15
    "%d-%b-%Y",       # 15-Jan-2024
    "%d %b %Y",       # 15 Jan 2024
    "%d %B %Y",       # 15 January 2024
    "%b %d, %Y",      # Jan 15, 2024
    "%B %d, %Y",      # January 15, 2024
    "%Y%m%d",         # 20240115
    "%d.%m.%Y",       # 15.01.2024
]


def _try_parse_date(value) -> Optional[date]:
    """Attempt to parse a date from various formats. Returns None on failure."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None

    # Already a date / datetime
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value

    raw = str(value).strip()
    if not raw:
        return None

    # Try ISO format first (YYYY-MM-DD) — unambiguous, no dayfirst confusion
    try:
        parsed = pd.to_datetime(raw, format="%Y-%m-%d")
        # Check if date is within Python's datetime range
        if parsed.year < 1 or parsed.year > 9999:
            return None
        return parsed.date()
    except (ValueError, TypeError):
        pass

    # Manual format attempts
    for fmt in _DATE_FORMATS:
        try:
            dt = datetime.strptime(raw, fmt)
            # Check if date is within Python's datetime range
            if dt.year < 1 or dt.year > 9999:
                return None
            return dt.date()
        except (ValueError, TypeError):
            continue

    # Last-resort pandas parsing for unusual formats
    for dayfirst in (True, False):
        try:
            parsed = pd.to_datetime(raw, dayfirst=dayfirst)
            # Check if date is within Python's datetime range
            if parsed.year < 1 or parsed.year > 9999:
                return None
            return parsed.date()
        except (ValueError, TypeError):
            continue

    return None


# ═══════════════════════════════════════════════════════════════════════════
# Amount parsing
# ═══════════════════════════════════════════════════════════════════════════

_CURRENCY_RE = re.compile(r"[^\d.\-+]")  # strip currency symbols / commas


def _try_parse_amount(value) -> Optional[float]:
    """Parse a monetary amount, stripping currency symbols and commas."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    if isinstance(value, (int, float)):
        return float(value)

    raw = str(value).strip()
    if not raw:
        return None

    # Handle parentheses-as-negative: (123.45) → -123.45
    if raw.startswith("(") and raw.endswith(")"):
        raw = "-" + raw[1:-1]

    cleaned = _CURRENCY_RE.sub("", raw)
    try:
        return float(cleaned)
    except (ValueError, TypeError):
        return None


# ═══════════════════════════════════════════════════════════════════════════
# Column detection
# ═══════════════════════════════════════════════════════════════════════════

def _normalize_col(name: str) -> str:
    """Lowercase, strip, collapse whitespace."""
    return re.sub(r"\s+", " ", str(name).strip().lower())


def _match_column(columns: List[str], keywords: List[str]) -> Optional[str]:
    """
    Find the best matching column name for a set of keywords.
    Returns the *original* column name (preserving case).
    """
    normalized = {_normalize_col(c): c for c in columns}

    # Exact match first
    for kw in keywords:
        if kw in normalized:
            return normalized[kw]

    # Token-aware substring / contains match
    for kw in keywords:
        for norm, original in normalized.items():
            tokens = re.split(r"[\s._/\-]+", norm)
            if len(kw) <= 2:
                if kw in tokens or norm == kw:
                    return original
            elif kw in norm or norm in kw:
                return original

    return None


def _cell_matches_keyword(cell: str, keyword: str) -> bool:
    """Check if a normalized header cell matches a normalized keyword."""
    tokens = re.split(r"[\s._/\-]+", cell)
    if len(keyword) <= 2:
        return keyword in tokens or cell == keyword
    return keyword == cell or keyword in cell or cell in keyword


def _dedupe_column_names(columns: List[str]) -> List[str]:
    """Preserve the first header label and suffix later duplicates."""
    seen: dict[str, int] = {}
    deduped: List[str] = []

    for idx, raw_name in enumerate(columns):
        name = str(raw_name).strip() or f"Unnamed: {idx}"
        count = seen.get(name, 0)
        deduped.append(name if count == 0 else f"{name}.{count}")
        seen[name] = count + 1

    return deduped


def _score_header_candidate(values: List[object]) -> float:
    """Heuristic score for picking the real header row in noisy files."""
    normalized = [_normalize_col(v) for v in values if str(v).strip()]
    if len(normalized) < 2:
        return -1.0

    keyword_hits = 0
    alpha_cells = 0
    numeric_like = 0
    long_cells = 0
    date_like = 0
    amount_like = 0

    for cell in normalized:
        if any(_cell_matches_keyword(cell, kw) for kw in _HEADER_KEYWORDS):
            keyword_hits += 1
        if re.search(r"[a-zA-Z]", cell):
            alpha_cells += 1
        if re.fullmatch(r"[\d.\-+/]+", cell):
            numeric_like += 1
        if len(cell) > 40:
            long_cells += 1
        if _try_parse_date(cell) is not None:
            date_like += 1
        if _try_parse_amount(cell) is not None:
            amount_like += 1

    unique_ratio = len(set(normalized)) / len(normalized)

    return (
        keyword_hits * 4.0
        + alpha_cells * 0.5
        + unique_ratio * 1.5
        - numeric_like * 1.0
        - long_cells * 1.5
        - date_like * 2.0
        - amount_like * 1.5
    )


def _count_header_hits(values: List[object]) -> int:
    """Count how many cells look like header labels rather than row data."""
    normalized = [_normalize_col(v) for v in values if str(v).strip()]
    hits = 0

    for cell in normalized:
        if any(_cell_matches_keyword(cell, kw) for kw in _HEADER_KEYWORDS):
            hits += 1

    return hits


def _promote_header_row(raw_df: pd.DataFrame, source_label: str) -> pd.DataFrame:
    """Choose the most likely header row, then promote it to dataframe columns."""
    if raw_df.empty:
        return raw_df

    search_window = min(len(raw_df), 10)
    best_row_index = 0
    best_score = float("-inf")
    best_hits = 0

    for idx in range(search_window):
        row_values = raw_df.iloc[idx].tolist()
        score = _score_header_candidate(row_values)
        hits = _count_header_hits(row_values)
        if score > best_score:
            best_score = score
            best_row_index = idx
            best_hits = hits

    if best_hits < _MIN_HEADER_HITS:
        columns = _dedupe_column_names(
            [f"column_{i}" for i in range(raw_df.shape[1])]
        )
        logger.info(
            "No confident header row for %s; using generated columns=%s",
            source_label,
            columns,
        )
        df = raw_df.copy()
        df.columns = columns
        return df

    header_values = [
        "" if pd.isna(value) else str(value).strip()
        for value in raw_df.iloc[best_row_index].tolist()
    ]
    columns = _dedupe_column_names(header_values)

    logger.info(
        "Detected header row for %s at index=%d columns=%s",
        source_label,
        best_row_index,
        columns,
    )

    df = raw_df.iloc[best_row_index + 1:].copy()
    df.columns = columns
    return df


def _cleanup_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """Trim whitespace, drop empty rows/columns, and reset the row index."""
    if df.empty:
        return df

    df = df.apply(
        lambda series: series.map(lambda value: value.strip() if isinstance(value, str) else value)
    )
    df.replace(r"^\s*$", pd.NA, regex=True, inplace=True)

    keep_columns = [
        col for col in df.columns
        if not df[col].isna().all()
    ]
    df = df.loc[:, keep_columns]

    df.columns = [str(col).strip() for col in df.columns]
    df.dropna(how="all", inplace=True)
    df.reset_index(drop=True, inplace=True)
    return df


def _read_delimited_file(path: Path) -> pd.DataFrame:
    """Read CSV/TSV/TXT data while auto-detecting the delimiter."""
    rows: Optional[List[List[str]]] = None
    selected_encoding: Optional[str] = None
    default_delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    selected_delimiter = default_delimiter

    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            with open(path, newline="", encoding=encoding) as handle:
                sample = handle.read(4096)
                handle.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
                    selected_delimiter = dialect.delimiter
                    rows = list(csv.reader(handle, dialect))
                except csv.Error:
                    selected_delimiter = default_delimiter
                    rows = list(csv.reader(handle, delimiter=default_delimiter))
            selected_encoding = encoding
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        raise ValueError("Could not decode delimited text file with any known encoding")

    logger.info(
        "Read delimited file %s with encoding=%s delimiter=%s raw_rows=%d",
        path.name,
        selected_encoding,
        selected_delimiter,
        len(rows or []),
    )

    if not rows:
        return pd.DataFrame()

    max_width = max(len(row) for row in rows)
    padded_rows = [
        row + [""] * (max_width - len(row))
        for row in rows
    ]
    return pd.DataFrame(padded_rows)


def _read_pdf_file(path: Path) -> pd.DataFrame:
    """Extract tabular data from PDF files using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("PDF processing requires 'pdfplumber' package. Run 'pip install pdfplumber'.")

    all_tables = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if table:
                    all_tables.extend(table)

    if not all_tables:
        # Fallback: extract text and try to parse as CSV-like
        text_data = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines = text.split('\n')
                    for line in lines:
                        # Try to split on common delimiters
                        if '\t' in line:
                            row = line.split('\t')
                        elif ',' in line:
                            row = line.split(',')
                        elif '|' in line:
                            row = line.split('|')
                        else:
                            # Assume space-separated, but this is tricky
                            row = line.split()
                        if row:
                            text_data.append(row)
        all_tables = text_data

    if not all_tables:
        raise ValueError("No tabular data found in PDF file.")

    # Find max width
    max_width = max(len(row) for row in all_tables) if all_tables else 0
    padded_rows = [row + [""] * (max_width - len(row)) for row in all_tables]
    return pd.DataFrame(padded_rows)


def _read_word_file(path: Path) -> pd.DataFrame:
    """Extract tabular data from Word documents."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("Word processing requires 'python-docx' package. Run 'pip install python-docx'.")

    doc = Document(path)
    all_tables = []

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            all_tables.append(row_data)

    if not all_tables:
        # Fallback: extract text and try to parse
        text = '\n'.join([para.text for para in doc.paragraphs])
        lines = text.split('\n')
        text_data = []
        for line in lines:
            if line.strip():
                # Try delimiters
                if '\t' in line:
                    row = line.split('\t')
                elif ',' in line:
                    row = line.split(',')
                elif '|' in line:
                    row = line.split('|')
                else:
                    row = line.split()
                text_data.append(row)
        all_tables = text_data

    if not all_tables:
        raise ValueError("No tabular data found in Word document.")

    max_width = max(len(row) for row in all_tables) if all_tables else 0
    padded_rows = [row + [""] * (max_width - len(row)) for row in all_tables]
    return pd.DataFrame(padded_rows)


def _read_plain_text_file(path: Path) -> pd.DataFrame:
    """Read plain text files and attempt to parse as tabular data."""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return pd.DataFrame()

    # Try to detect delimiter
    sample_line = lines[0] if lines else ""
    if '\t' in sample_line:
        delimiter = '\t'
    elif ',' in sample_line:
        delimiter = ','
    elif '|' in sample_line:
        delimiter = '|'
    elif ';' in sample_line:
        delimiter = ';'
    else:
        # Assume space-separated, but limit splits
        delimiter = None

    parsed_rows = []
    for line in lines:
        if delimiter:
            row = line.split(delimiter)
        else:
            # Space separated, but keep multiple spaces as one
            row = line.split()
        parsed_rows.append(row)

    max_width = max(len(row) for row in parsed_rows) if parsed_rows else 0
    padded_rows = [row + [""] * (max_width - len(row)) for row in parsed_rows]
    return pd.DataFrame(padded_rows)


def _read_spreadsheet_file(path: Path) -> pd.DataFrame:
    """Read the most transaction-like worksheet from an Excel-family file."""
    ext = path.suffix.lower()
    engine = "openpyxl"

    if ext == ".xls":
        if importlib.util.find_spec("xlrd") is None:
            raise ValueError(
                "Legacy .xls uploads require the 'xlrd' package. "
                "Run 'pip install -r server/requirements.txt' to enable .xls parsing."
            )
        engine = "xlrd"

    sheets = pd.read_excel(path, header=None, sheet_name=None, engine=engine)
    if not sheets:
        return pd.DataFrame()

    best_sheet_name = None
    best_sheet = None
    best_score = float("-inf")

    for sheet_name, raw_df in sheets.items():
        if raw_df.empty:
            continue

        search_window = min(len(raw_df), 10)
        header_score = max(
            _score_header_candidate(raw_df.iloc[idx].tolist())
            for idx in range(search_window)
        )
        non_empty_rows = raw_df.dropna(how="all").shape[0]
        non_empty_cells = int(raw_df.notna().sum().sum())
        score = header_score + (non_empty_rows * 0.1) + (non_empty_cells * 0.01)

        if score > best_score:
            best_score = score
            best_sheet_name = sheet_name
            best_sheet = raw_df

    if best_sheet is None:
        first_sheet_name, first_sheet = next(iter(sheets.items()))
        best_sheet_name = first_sheet_name
        best_sheet = first_sheet

    logger.info(
        "Selected spreadsheet sheet for %s: %s score=%.2f",
        path.name,
        best_sheet_name,
        best_score,
    )
    return best_sheet


def _detect_date_by_content(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    """
    If header-based detection fails, try parsing the first few values
    of each column to find one that looks like dates.
    """
    sample_size = min(5, len(df))
    for col in candidates:
        sample = df[col].head(sample_size)
        parsed_count = sum(1 for v in sample if _try_parse_date(v) is not None)
        if parsed_count >= sample_size * 0.6:
            return col
    return None


def _detect_amount_by_content(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    """
    If header-based detection fails, try parsing the first few values
    of each column to find one that looks like amounts.
    """
    sample_size = min(5, len(df))
    for col in candidates:
        sample = df[col].head(sample_size)
        parsed_count = sum(1 for v in sample if _try_parse_amount(v) is not None)
        if parsed_count >= sample_size * 0.6:
            return col
    return None


def _detect_type_by_content(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    """
    Detect a type column by looking for values like debit/credit/dr/cr.
    """
    sample_size = min(5, len(df))
    for col in candidates:
        sample = df[col].head(sample_size)
        parsed_count = sum(
            1 for v in sample
            if _determine_type_from_column(v) != TransactionType.UNKNOWN
        )
        if parsed_count >= max(1, int(sample_size * 0.6)):
            return col
    return None


def _looks_like_credit_card_layout(columns: List[object]) -> bool:
    """
    Credit-card statements often have both transaction and post dates with
    positive charges and negative payments/credits.
    """
    normalized = {_normalize_col(c) for c in columns}
    return any("post date" in col or "posting date" in col for col in normalized)


def detect_columns(df: pd.DataFrame) -> ColumnMapping:
    """
    Detect which columns correspond to date, description, amount,
    debit, credit, and type using header keywords + content probing.
    """
    cols = list(df.columns)
    matched: set = set()
    confidence_hits = 0
    total_checks = 4  # date, description, amount, debit/credit

    # --- Date ---------------------------------------------------------------
    date_col = _match_column(cols, _DATE_KEYWORDS)
    if date_col is None:
        remaining = [c for c in cols if c not in matched]
        date_col = _detect_date_by_content(df, remaining)
    if date_col:
        matched.add(date_col)
        confidence_hits += 1

    # --- Description --------------------------------------------------------
    desc_col = _match_column(cols, _DESCRIPTION_KEYWORDS)
    if desc_col and desc_col not in matched:
        matched.add(desc_col)
        confidence_hits += 1
    elif desc_col in matched:
        desc_col = None

    # If still no description, pick the column with the longest average string
    if desc_col is None:
        remaining_str = [
            c for c in cols
            if c not in matched and df[c].dtype == object
        ]
        if remaining_str:
            avg_lens = {
                c: df[c].astype(str).str.len().mean() for c in remaining_str
            }
            best = max(avg_lens, key=avg_lens.get)
            desc_col = best
            matched.add(desc_col)
            confidence_hits += 0.5  # lower confidence for guessed match

    # --- Amount (single column) ---------------------------------------------
    amount_col = _match_column(
        [c for c in cols if c not in matched], _AMOUNT_KEYWORDS
    )
    if amount_col is None:
        remaining = [c for c in cols if c not in matched]
        amount_col = _detect_amount_by_content(df, remaining)
    if amount_col:
        matched.add(amount_col)
        confidence_hits += 1

    # --- Debit / Credit (separate columns) ----------------------------------
    debit_col = _match_column(
        [c for c in cols if c not in matched], _DEBIT_KEYWORDS
    )
    credit_col = _match_column(
        [c for c in cols if c not in matched], _CREDIT_KEYWORDS
    )
    if debit_col:
        matched.add(debit_col)
    if credit_col:
        matched.add(credit_col)

    if debit_col or credit_col:
        confidence_hits += 1

    # --- Type column --------------------------------------------------------
    type_col = _match_column(
        [c for c in cols if c not in matched], _TYPE_KEYWORDS
    )
    if type_col is None:
        remaining = [c for c in cols if c not in matched]
        type_col = _detect_type_by_content(df, remaining)
    if type_col:
        matched.add(type_col)

    confidence = round(min(confidence_hits / total_checks, 1.0), 2)

    return ColumnMapping(
        date_column=date_col,
        description_column=desc_col,
        amount_column=amount_col,
        debit_column=debit_col,
        credit_column=credit_col,
        type_column=type_col,
        confidence=confidence,
    )


# ═══════════════════════════════════════════════════════════════════════════
# Row normalization
# ═══════════════════════════════════════════════════════════════════════════

def _clean_description(value) -> str:
    """Normalize a description string."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    text = str(value).strip()
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text)
    return text


def _determine_type_from_column(value) -> TransactionType:
    """Infer transaction type from a 'type' column value."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return TransactionType.UNKNOWN
    v = str(value).strip().lower()
    if v in ("debit", "dr", "withdrawal", "deduction", "spent"):
        return TransactionType.DEBIT
    if v in ("credit", "cr", "deposit", "addition", "received"):
        return TransactionType.CREDIT
    return TransactionType.UNKNOWN


def _normalize_row(
    row: pd.Series,
    row_index: int,
    mapping: ColumnMapping,
) -> Transaction:
    """Convert a single DataFrame row into a normalized Transaction."""

    notes: List[str] = []
    original = row.to_dict()

    # ── Date ────────────────────────────────────────────────────────────
    parsed_date = None
    if mapping.date_column and mapping.date_column in row.index:
        parsed_date = _try_parse_date(row[mapping.date_column])
        if parsed_date is None:
            notes.append(f"Could not parse date: {row[mapping.date_column]}")

    # ── Description ─────────────────────────────────────────────────────
    description = ""
    if mapping.description_column and mapping.description_column in row.index:
        description = _clean_description(row[mapping.description_column])

    # ── Amount & Type ───────────────────────────────────────────────────
    amount = 0.0
    amount_parsed = False
    txn_type = TransactionType.UNKNOWN

    # Case 1: Separate debit/credit columns
    if mapping.debit_column or mapping.credit_column:
        debit_val = None
        credit_val = None

        if mapping.debit_column and mapping.debit_column in row.index:
            debit_val = _try_parse_amount(row[mapping.debit_column])

        if mapping.credit_column and mapping.credit_column in row.index:
            credit_val = _try_parse_amount(row[mapping.credit_column])

        if debit_val is not None:
            amount_parsed = True
        if credit_val is not None:
            amount_parsed = True

        if debit_val is not None and debit_val != 0:
            amount = abs(debit_val)
            txn_type = TransactionType.DEBIT
        elif credit_val is not None and credit_val != 0:
            amount = abs(credit_val)
            txn_type = TransactionType.CREDIT
        elif mapping.amount_column and mapping.amount_column in row.index:
            # Fallback to amount column if debit/credit are empty
            parsed_amount = _try_parse_amount(row[mapping.amount_column])
            if parsed_amount is not None:
                amount_parsed = True
                amount = abs(parsed_amount)
                txn_type = (
                    TransactionType.DEBIT if parsed_amount < 0
                    else TransactionType.CREDIT if parsed_amount > 0
                    else TransactionType.UNKNOWN
                )

    # Case 2: Single amount column
    elif mapping.amount_column and mapping.amount_column in row.index:
        parsed_amount = _try_parse_amount(row[mapping.amount_column])
        if parsed_amount is not None:
            amount_parsed = True
            amount = abs(parsed_amount)
            is_credit_card_layout = _looks_like_credit_card_layout(list(row.index))
            # Negative → debit, positive → credit (common convention)
            txn_type = (
                TransactionType.CREDIT if is_credit_card_layout and parsed_amount < 0
                else TransactionType.DEBIT if is_credit_card_layout and parsed_amount > 0
                else TransactionType.DEBIT if parsed_amount < 0
                else TransactionType.CREDIT if parsed_amount > 0
                else TransactionType.UNKNOWN
            )
        else:
            notes.append(
                f"Could not parse amount: {row[mapping.amount_column]}"
            )

    # Override type if a dedicated type column exists
    if mapping.type_column and mapping.type_column in row.index:
        detected = _determine_type_from_column(row[mapping.type_column])
        if detected != TransactionType.UNKNOWN:
            txn_type = detected

    # ── Validation ──────────────────────────────────────────────────────
    is_valid = True
    if parsed_date is None:
        is_valid = False
        if not any("date" in n.lower() for n in notes):
            notes.append("Missing or unparseable date")
    if not amount_parsed:
        is_valid = False
        if not any("amount" in n.lower() for n in notes):
            notes.append("Missing or unparseable amount")
    if txn_type == TransactionType.UNKNOWN:
        is_valid = False
        notes.append("Missing or unrecognized transaction type")
    elif amount == 0.0:
        notes.append("Amount is zero")

    # Sanitize original_data — convert NaN to None for JSON
    clean_original = {
        k: (None if isinstance(v, float) and pd.isna(v) else v)
        for k, v in original.items()
    }

    return Transaction(
        row_index=row_index,
        transaction_date=parsed_date,
        description=description,
        amount=round(amount, 2),
        transaction_type=txn_type,
        original_data=clean_original,
        is_valid=is_valid,
        validation_notes=notes,
    )


# ═══════════════════════════════════════════════════════════════════════════
# File reading
# ═══════════════════════════════════════════════════════════════════════════

def read_file_to_dataframe(file_path: str) -> pd.DataFrame:
    """
    Read a supported file into a pandas DataFrame.
    Handles delimited text, spreadsheets, PDFs, Word docs, and plain text.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext in {".csv", ".tsv"}:
        raw_df = _read_delimited_file(path)
    elif ext in _SPREADSHEET_EXTENSIONS:
        raw_df = _read_spreadsheet_file(path)
    elif ext == ".pdf":
        raw_df = _read_pdf_file(path)
    elif ext in {".docx", ".doc"}:
        raw_df = _read_word_file(path)
    elif ext == ".txt":
        raw_df = _read_plain_text_file(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    df = _promote_header_row(raw_df, path.name)
    return _cleanup_dataframe(df)


# ═══════════════════════════════════════════════════════════════════════════
# Main parse function
# ═══════════════════════════════════════════════════════════════════════════

def parse_transactions(
    file_path: str,
    file_id: str,
    original_filename: str,
) -> ParseResult:
    """
    Full parse pipeline:
      1. Read file → DataFrame
      2. Detect columns
      3. Normalize each row
      4. Build summary statistics

    Returns a ParseResult with all transactions + metadata.
    """

    # Step 1: Read
    df = read_file_to_dataframe(file_path)

    logger.info(
        "Parser dataframe file_id=%s file=%s columns=%s normalized_columns=%s rows=%d",
        file_id,
        original_filename,
        list(df.columns),
        [_normalize_col(col) for col in df.columns],
        len(df),
    )
    if not df.empty:
        logger.info("Parser dataframe head file_id=%s\n%s", file_id, df.head().to_string())

    if df.empty:
        logger.warning("Parser produced an empty dataframe for file_id=%s file=%s", file_id, original_filename)
        return ParseResult(
            file_id=file_id,
            original_filename=original_filename,
            column_mapping=ColumnMapping(confidence=0.0),
            summary=ParseSummary(),
            transactions=[],
        )

    # Step 2: Detect columns
    mapping = detect_columns(df)
    logger.info("Parser detected columns file_id=%s mapping=%s", file_id, mapping.model_dump())

    # Step 3: Normalize rows
    transactions: List[Transaction] = []
    for idx, row in df.iterrows():
        txn = _normalize_row(row, row_index=int(idx) + 1, mapping=mapping)
        transactions.append(txn)

    # Step 4: Build summary
    valid = [t for t in transactions if t.is_valid]
    invalid = [t for t in transactions if not t.is_valid]

    valid_dates = [t.transaction_date for t in valid if t.transaction_date is not None]
    total_debits = round(
        sum(t.amount for t in valid if t.transaction_type == TransactionType.DEBIT), 2
    )
    total_credits = round(
        sum(t.amount for t in valid if t.transaction_type == TransactionType.CREDIT), 2
    )

    logger.info(
        "Parser structured transactions file_id=%s valid=%d invalid=%d sample=%s",
        file_id,
        len(valid),
        len(invalid),
        [txn.model_dump(mode="json", by_alias=True) for txn in valid[:5]],
    )

    summary = ParseSummary(
        total_rows=len(transactions),
        valid_rows=len(valid),
        invalid_rows=len(invalid),
        skipped_rows=0,
        date_range_start=min(valid_dates) if valid_dates else None,
        date_range_end=max(valid_dates) if valid_dates else None,
        total_debits=total_debits,
        total_credits=total_credits,
        net_amount=round(total_credits - total_debits, 2),
    )

    return ParseResult(
        file_id=file_id,
        original_filename=original_filename,
        column_mapping=mapping,
        summary=summary,
        transactions=transactions,
    )


def parse_transactions_with_mapping(
    file_path: str,
    file_id: str,
    original_filename: str,
    column_mapping: ColumnMapping,
) -> ParseResult:
    """
    Parse transactions using provided column mappings.
    Similar to parse_transactions but skips auto-detection.
    """

    # Step 1: Read
    df = read_file_to_dataframe(file_path)

    logger.info(
        "Parser dataframe file_id=%s file=%s columns=%s normalized_columns=%s rows=%d",
        file_id,
        original_filename,
        list(df.columns),
        [_normalize_col(col) for col in df.columns],
        len(df),
    )
    if not df.empty:
        logger.info("Parser dataframe head file_id=%s\n%s", file_id, df.head().to_string())

    if df.empty:
        return ParseResult(
            file_id=file_id,
            original_filename=original_filename,
            column_mapping=column_mapping,
            summary=ParseSummary(),
            transactions=[],
        )

    # Step 2: Use provided mapping
    mapping = column_mapping
    logger.info("Parser using provided columns file_id=%s mapping=%s", file_id, mapping.model_dump())

    # Step 3: Normalize rows
    transactions: List[Transaction] = []
    for idx, row in df.iterrows():
        txn = _normalize_row(row, row_index=int(idx) + 1, mapping=mapping)
        transactions.append(txn)

    # Step 4: Build summary
    valid = [t for t in transactions if t.is_valid]
    invalid = [t for t in transactions if not t.is_valid]

    valid_dates = [t.transaction_date for t in valid if t.transaction_date is not None]
    total_debits = round(
        sum(t.amount for t in valid if t.transaction_type == TransactionType.DEBIT), 2
    )
    total_credits = round(
        sum(t.amount for t in valid if t.transaction_type == TransactionType.CREDIT), 2
    )

    summary = ParseSummary(
        total_rows=len(transactions),
        valid_rows=len(valid),
        invalid_rows=len(invalid),
        skipped_rows=0,
        date_range_start=min(valid_dates) if valid_dates else None,
        date_range_end=max(valid_dates) if valid_dates else None,
        total_debits=total_debits,
        total_credits=total_credits,
        net_amount=round(total_credits - total_debits, 2),
    )

    return ParseResult(
        file_id=file_id,
        original_filename=original_filename,
        column_mapping=mapping,
        summary=summary,
        transactions=transactions,
    )


def validate_parse_result(parse_result: ParseResult) -> None:
    """Raise when a file cannot be interpreted as usable transaction data."""
    mapping = parse_result.column_mapping
    missing_columns: List[str] = []

    if not mapping.date_column:
        missing_columns.append("date")
    if not mapping.description_column:
        missing_columns.append("description")
    if not (mapping.amount_column or mapping.debit_column or mapping.credit_column):
        missing_columns.append("amount")

    if missing_columns:
        raise ValueError(
            "Could not detect the required transaction columns: "
            + ", ".join(missing_columns)
        )

    if parse_result.summary.valid_rows == 0:
        raise ValueError(
            "No valid transactions found after parsing. Check that the file contains "
            "date, description, amount, and type/debit/credit columns."
        )
